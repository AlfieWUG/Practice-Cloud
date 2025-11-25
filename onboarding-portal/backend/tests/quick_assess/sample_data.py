"""Utilities to generate sample Quick Assess documents for tests."""
from __future__ import annotations

from io import BytesIO
from zipfile import ZipFile

from docx import Document
from reportlab.pdfgen import canvas


def sample_docx_bytes() -> bytes:
    """Create a simple DOCX document in memory."""
    document = Document()
    document.add_heading("Infrastructure Design", level=1)
    document.add_paragraph("Web tier: 2 x nginx servers running Ubuntu 22.04.")
    document.add_paragraph("App tier: Java Spring services.")
    document.add_paragraph("Data tier: PostgreSQL 14 primary / replica.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def sample_pdf_bytes() -> bytes:
    """Create a basic PDF document in memory."""
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.setTitle("Sample Infrastructure PDF")
    pdf.drawString(72, 750, "Infrastructure Summary")
    pdf.drawString(72, 730, "Load balancer -> Web -> App -> Database")
    pdf.drawString(72, 710, "Technologies: AWS, Docker, PostgreSQL")
    pdf.save()
    return buffer.getvalue()


def sample_drawio_bytes() -> bytes:
    """Return a simple draw.io XML diagram."""
    xml = """
    <mxfile host="app.diagrams.net">
      <diagram name="Data Flow">
        <mxGraphModel>
          <root>
            <mxCell id="0"/>
            <mxCell id="1" parent="0"/>
            <mxCell id="2" value="Web Server" vertex="1" parent="1">
              <mxGeometry x="100" y="100" width="80" height="30" as="geometry"/>
            </mxCell>
            <mxCell id="3" value="App Server" vertex="1" parent="1">
              <mxGeometry x="250" y="100" width="80" height="30" as="geometry"/>
            </mxCell>
            <mxCell id="4" value="Database" vertex="1" parent="1">
              <mxGeometry x="400" y="100" width="80" height="30" as="geometry"/>
            </mxCell>
            <mxCell id="5" edge="1" source="2" target="3" parent="1">
              <mxGeometry relative="1" as="geometry"/>
            </mxCell>
            <mxCell id="6" edge="1" source="3" target="4" parent="1">
              <mxGeometry relative="1" as="geometry"/>
            </mxCell>
          </root>
        </mxGraphModel>
      </diagram>
    </mxfile>
    """.strip()
    return xml.encode("utf-8")


def sample_vsdx_bytes() -> bytes:
    """Create a minimal Visio .vsdx archive with two shapes and a connector."""
    page_xml = """
    <VisioDocument xmlns:v="http://schemas.microsoft.com/office/visio/2012/main">
      <Pages>
        <PageContents>
          <Shapes>
            <Shape ID="1">
              <Text><cp>Web Tier</cp></Text>
            </Shape>
            <Shape ID="2">
              <Text><cp>Data Tier</cp></Text>
            </Shape>
          </Shapes>
          <Connects>
            <Connect FromSheet="1" ToSheet="2"/>
          </Connects>
        </PageContents>
      </Pages>
    </VisioDocument>
    """.strip()

    buffer = BytesIO()
    with ZipFile(buffer, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            """<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
            <Default Extension="xml" ContentType="application/xml"/>
            </Types>""",
        )
        archive.writestr("visio/pages/page1.xml", page_xml)
    return buffer.getvalue()

