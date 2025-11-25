"""
DocumentParserAgent implemented with LangGraph to process Word/PDF files.

Capabilities:
- Extract text and structural metadata from DOCX/PDF
- Detect sections, tables, and document metadata
- Run Claude Sonnet 4.5 via AWS Bedrock for technical entity extraction
- Provide structured JSON output compatible with Quick Assess workflows
"""

from __future__ import annotations

import asyncio
import io
import json
import logging
import re
from dataclasses import dataclass
from typing import Any, Dict, List, Literal, Optional, TypedDict

import pdfplumber
from docx import Document as DocxDocument
from langgraph.graph import END, StateGraph

from agentic_services.agents.base import BaseAgent
from agentic_services.config.settings import settings

logger = logging.getLogger(__name__)


ALLOWED_TYPES = {
    ".docx": "word",
    ".pdf": "pdf",
}


class DocumentParserState(TypedDict, total=False):
    """State shared across LangGraph nodes."""

    document_name: str
    document_type: Literal["word", "pdf"]
    source_path: Optional[str]
    s3_uri: Optional[str]
    file_bytes: Optional[bytes]
    raw_bytes: bytes
    extracted_text: str
    structure: Dict[str, Any]
    metadata: Dict[str, Any]
    technical_entities: Dict[str, Any]
    heuristic_entities: Dict[str, List[str]]
    errors: List[str]


class DocumentParseError(Exception):
    """Raised when a document cannot be parsed."""


@dataclass
class ParsedDocument:
    """Convenience container."""

    text: str
    structure: Dict[str, Any]
    metadata: Dict[str, Any]


class DocumentParserAgent(BaseAgent):
    """LangGraph-driven agent for document parsing."""

    SYSTEM_PROMPT = """You are DocumentParserAgent inside Nagarro's AIMS platform.
Given raw document text, identify:
- Server or host names
- IP addresses
- Technologies or platforms mentioned

Return STRICT JSON:
{
  "servers": [list of strings],
  "ip_addresses": [list of IPv4/IPv6 strings],
  "technologies": [list of technologies, vendors, frameworks]
}

Do not add commentary."""

    def __init__(self, agent_id: Optional[str] = None):
        super().__init__(agent_id=agent_id)
        self.graph = self._build_graph()

    def _build_graph(self):
        graph = StateGraph(DocumentParserState)
        graph.add_node("load_document", self._load_document)
        graph.add_node("extract_content", self._extract_content)
        graph.add_node("identify_entities", self._identify_entities)
        graph.add_node("aggregate", self._aggregate_results)

        graph.set_entry_point("load_document")
        graph.add_edge("load_document", "extract_content")
        graph.add_edge("extract_content", "identify_entities")
        graph.add_edge("identify_entities", "aggregate")
        graph.add_edge("aggregate", END)
        return graph.compile()

    async def execute(self, task: Dict[str, Any]) -> Dict[str, Any]:
        """Execute the LangGraph pipeline and return structured JSON."""
        self.validate_task(task, ["document_name"])
        if not any([task.get("file_path"), task.get("s3_uri"), task.get("file_bytes")]):
            raise ValueError("One of file_path, s3_uri, or file_bytes must be provided")

        logger.info(
            "DocumentParserAgent started for %s",
            task["document_name"],
        )

        initial_state: DocumentParserState = {
            "document_name": task["document_name"],
            "source_path": task.get("file_path"),
            "s3_uri": task.get("s3_uri"),
            "file_bytes": task.get("file_bytes"),
            "errors": [],
        }

        try:
            result_state = await self.graph.ainvoke(initial_state)
        except Exception as exc:
            logger.exception("DocumentParserAgent failed: %s", exc)
            raise

        output = {
            "document_name": task["document_name"],
            "document_type": result_state.get("document_type"),
            "extracted_text": result_state.get("extracted_text", ""),
            "structure": result_state.get("structure", {}),
            "metadata": result_state.get("metadata", {}),
            "technical_entities": result_state.get("technical_entities", {}),
        }
        if result_state.get("errors"):
            output["errors"] = result_state["errors"]
        return output

    # ----- LangGraph nodes -----
    async def _load_document(self, state: DocumentParserState) -> DocumentParserState:
        """Load bytes from local path, S3 URI, or direct bytes."""
        document_name = state["document_name"]
        document_type = self._infer_document_type(document_name)

        if document_type not in ("word", "pdf"):
            raise DocumentParseError(
                f"Unsupported document type for {document_name}. Allowed: {list(ALLOWED_TYPES.values())}"
            )

        if state.get("file_bytes"):
            raw_bytes = state["file_bytes"]  # type: ignore[assignment]
        elif state.get("source_path"):
            raw_bytes = await asyncio.to_thread(
                self._read_local_file,
                state["source_path"],
            )
        elif state.get("s3_uri"):
            raw_bytes = await self._read_s3_file(state["s3_uri"])
        else:
            raise DocumentParseError("No document source provided")

        return {
            "raw_bytes": raw_bytes,
            "document_type": document_type,
        }

    async def _extract_content(self, state: DocumentParserState) -> DocumentParserState:
        """Extract text + structure from DOCX/PDF."""
        raw_bytes = state["raw_bytes"]
        document_type = state["document_type"]

        try:
            if document_type == "word":
                parsed = await asyncio.to_thread(self._parse_docx, raw_bytes)
            else:
                parsed = await asyncio.to_thread(self._parse_pdf, raw_bytes)
        except Exception as exc:
            raise DocumentParseError(f"Failed to extract content: {exc}") from exc

        heuristic_entities = self._heuristic_entities(parsed.text)

        return {
            "extracted_text": parsed.text,
            "structure": parsed.structure,
            "metadata": parsed.metadata,
            "heuristic_entities": heuristic_entities,
        }

    async def _identify_entities(self, state: DocumentParserState) -> DocumentParserState:
        """Invoke Claude Sonnet 4.5 via Bedrock to enrich entity extraction."""
        text = state.get("extracted_text", "")
        if not text.strip():
            return {"technical_entities": self._combine_entities(state, {})}

        truncated = text[:12000]
        prompt = (
            "Analyze the following document text and extract infrastructure-related entities.\n\n"
            f"TEXT:\n{truncated}\n"
        )

        try:
            response_text = await self._invoke_bedrock(prompt)
            llm_entities = self._safe_parse_json(response_text)
        except Exception as exc:
            logger.warning("LLM entity extraction failed: %s", exc)
            llm_entities = {}

        combined = self._combine_entities(state, llm_entities)
        return {"technical_entities": combined}

    async def _aggregate_results(self, state: DocumentParserState) -> DocumentParserState:
        """Placeholder node to keep LangGraph happy."""
        return state

    # ----- Helpers -----
    def _infer_document_type(self, filename: str) -> str:
        suffix = filename.lower().rsplit(".", 1)[-1]
        return "word" if suffix == "docx" else "pdf"

    def _read_local_file(self, path: str) -> bytes:
        with open(path, "rb") as file:
            return file.read()

    async def _read_s3_file(self, uri: str) -> bytes:
        bucket, key = self._parse_s3_uri(uri)
        response = await asyncio.to_thread(
            self.s3.client.get_object, Bucket=bucket, Key=key
        )
        return response["Body"].read()

    def _parse_s3_uri(self, uri: str) -> (str, str):
        if not uri.startswith("s3://"):
            raise ValueError("Invalid S3 URI")
        bucket_key = uri[len("s3://") :]
        bucket, _, key = bucket_key.partition("/")
        if not bucket or not key:
            raise ValueError("S3 URI must include bucket and key")
        return bucket, key

    def _parse_docx(self, raw_bytes: bytes) -> ParsedDocument:
        try:
            doc = DocxDocument(io.BytesIO(raw_bytes))
        except Exception as exc:
            raise DocumentParseError("Unable to open DOCX file") from exc

        text_lines: List[str] = []
        sections: List[str] = []
        for paragraph in doc.paragraphs:
            value = paragraph.text.strip()
            if not value:
                continue
            text_lines.append(value)
            if paragraph.style and "Heading" in paragraph.style.name:
                sections.append(value)

        tables_count = len(doc.tables)
        metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "created": doc.core_properties.created.isoformat()
            if doc.core_properties.created
            else None,
            "modified": doc.core_properties.modified.isoformat()
            if doc.core_properties.modified
            else None,
            "pages": None,
            "tables_count": tables_count,
        }

        structure = {
            "sections": sections,
            "tables_count": tables_count,
        }

        return ParsedDocument(
            text="\n".join(text_lines),
            structure=structure,
            metadata=metadata,
        )

    def _parse_pdf(self, raw_bytes: bytes) -> ParsedDocument:
        try:
            pdf = pdfplumber.open(io.BytesIO(raw_bytes))
        except Exception as exc:
            raise DocumentParseError("Unable to open PDF file") from exc

        text_lines: List[str] = []
        sections: List[str] = []
        total_tables = 0
        try:
            for page in pdf.pages:
                content = page.extract_text() or ""
                text_lines.append(content)
                inferred_sections = self._infer_sections_from_page(content)
                sections.extend(inferred_sections)
                total_tables += len(page.find_tables())
        finally:
            pdf.close()

        metadata_raw = pdf.metadata or {}
        metadata = {
            "title": metadata_raw.get("Title") or "",
            "author": metadata_raw.get("Author") or "",
            "pages": metadata_raw.get("Pages") or len(text_lines),
            "creation_date": metadata_raw.get("CreationDate"),
            "mod_date": metadata_raw.get("ModDate"),
        }

        structure = {
            "sections": sections,
            "tables_count": total_tables,
        }

        return ParsedDocument(
            text="\n".join(text_lines),
            structure=structure,
            metadata=metadata,
        )

    def _infer_sections_from_page(self, text: str) -> List[str]:
        sections = []
        for line in text.splitlines():
            stripped = line.strip()
            if len(stripped) < 4:
                continue
            if stripped.isupper() or stripped.endswith(":"):
                sections.append(stripped[:120])
        return sections

    def _heuristic_entities(self, text: str) -> Dict[str, List[str]]:
        servers = sorted({match.group(0) for match in re.finditer(r"\b[\w-]*srv[\w-]*\b", text, re.IGNORECASE)})
        ips = sorted({match.group(0) for match in re.finditer(r"\b(?:\d{1,3}\.){3}\d{1,3}\b", text)})
        tech_keywords = sorted(
            {
                token
                for token in re.findall(r"\b[A-Z][A-Za-z0-9\-\+]{2,}\b", text)
                if token.upper() not in {"THE", "AND", "FOR", "WITH"}
            }
        )
        return {
            "servers": servers,
            "ip_addresses": ips,
            "technologies": tech_keywords[:50],
        }

    async def _invoke_bedrock(self, prompt: str) -> str:
        response = await self.bedrock.invoke_claude(
            prompt=prompt,
            system_prompt=self.SYSTEM_PROMPT,
            temperature=0.1,
            max_tokens=2000,
            model_id=settings.BEDROCK_DOCUMENT_PARSER_MODEL_ID,
        )
        return response["text"]

    def _safe_parse_json(self, text: str) -> Dict[str, Any]:
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            match = re.search(r"\{.*\}", text, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group(0))
                except json.JSONDecodeError:
                    pass
        return {}

    def _combine_entities(
        self,
        state: DocumentParserState,
        llm_entities: Dict[str, Any],
    ) -> Dict[str, List[str]]:
        heuristic = state.get("heuristic_entities", {})

        def merge(key: str) -> List[str]:
            combined = set()
            for source in (heuristic, llm_entities):
                values = source.get(key, []) if isinstance(source, dict) else []
                for value in values:
                    if isinstance(value, str) and value.strip():
                        combined.add(value.strip())
            return sorted(combined)

        return {
            "servers": merge("servers"),
            "ip_addresses": merge("ip_addresses"),
            "technologies": merge("technologies"),
        }

